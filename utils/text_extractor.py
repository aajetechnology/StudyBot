import base64
import io
import os
from groq import Groq
from pypdf import PdfReader
import pypdfium2 as pdfium
import docx

client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

def call_groq_vision(base64_image, prompt_text):
    try:
        response = client.chat.completions.create(
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt_text},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                ]
            }],
            model="llama-3.2-11b-vision-preview",
        )
        return response.choices[0].message.content + "\n"
    except Exception as e:
        return f"\n[Vision Error: {str(e)}]\n"

def extract_text_from_file(file_obj_or_path):
    """
    Unified text extraction for PDF, DOCX, and Images.
    Supports both Flask file objects and local file paths.
    """
    if isinstance(file_obj_or_path, str):
        # It's a path
        filename = file_obj_or_path.lower()
        f = open(file_obj_or_path, 'rb')
    else:
        # It's a file object (e.g. from Flask request)
        filename = file_obj_or_path.filename.lower()
        f = file_obj_or_path

    extracted_text = ""
    
    try:
        if filename.endswith('.pdf'):
            f.seek(0)
            reader = PdfReader(f)
            for page in reader.pages:
                extracted_text += page.extract_text() or ""
            
            if len(extracted_text.strip()) < 50:
                f.seek(0)
                pdf = pdfium.PdfDocument(f)
                for i in range(min(5, len(pdf))):
                    page = pdf[i]
                    bitmap = page.render(scale=2)
                    pil_image = bitmap.to_pil()
                    img_byte_arr = io.BytesIO()
                    pil_image.save(img_byte_arr, format='JPEG')
                    base64_image = base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')
                    extracted_text += call_groq_vision(base64_image, "Extract lecture text.")
                pdf.close()
        elif filename.endswith(('.png', '.jpg', '.jpeg')):
            f.seek(0)
            image_bytes = f.read()
            base64_image = base64.b64encode(image_bytes).decode('utf-8')
            extracted_text = call_groq_vision(base64_image, "Transcribe perfectly.")
        elif filename.endswith('.docx'):
            f.seek(0)
            doc = docx.Document(f)
            full_text = []
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            extracted_text = "\n".join(full_text)
            if not extracted_text.strip():
                extracted_text = "[System Warning: Word document appeared to be empty or contains only non-text elements.]"
        elif filename.endswith('.txt'):
            f.seek(0)
            extracted_text = f.read().decode('utf-8', errors='ignore')
    except Exception as e:
        print(f"Extraction Error for {filename}: {e}")
        extracted_text = f"\n[System Error reading {filename}: {str(e)}]\n"
    finally:
        if isinstance(file_obj_or_path, str):
            f.close()
            
    return extracted_text

import os
from docx import Document
from xhtml2pdf import pisa
import markdown

def save_study_notes(summary, transcript, output_path):
    root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    final_destination = os.path.join(root_dir, 'output', os.path.basename(output_path))
    os.makedirs(os.path.dirname(final_destination), exist_ok=True)

    if final_destination.endswith('.docx'):
        doc = Document()
        doc.add_heading('Professional Study Guide', 0)
        
        # Improved formatting: Parses Markdown-style lines from AI
        for line in summary.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], 2)
            elif line.startswith('- '):
                doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        doc.add_page_break()
        doc.add_heading('Transcript', 1)
        doc.add_paragraph(transcript)
        doc.save(final_destination)
    else:
        # PDF Generation with Professional CSS
        html_body = markdown.markdown(summary)
        html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: 'Helvetica', 'Arial', sans-serif; line-height: 1.6; padding: 40px; color: #222; }}
                h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 5px; }}
                h2 {{ color: #2980b9; margin-top: 25px; }}
                p {{ margin-bottom: 12px; }}
                .transcript {{ color: #666; font-size: 0.85em; border-top: 1px solid #eee; padding-top: 20px; }}
            </style>
        </head>
        <body>
            {html_body}
            <div class='transcript'>
                <h1>Raw Transcript</h1>
                <p>{transcript.replace('\n', '<br>')}</p>
            </div>
        </body>
        </html>
        """
        with open(final_destination, "wb") as f:
            pisa.CreatePDF(html, dest=f)
